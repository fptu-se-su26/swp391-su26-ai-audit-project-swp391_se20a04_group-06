# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, color):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = color
        run.font.bold = True
    return heading

def main():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette
    navy = RGBColor(11, 79, 108)      # #0B4F6C
    teal = RGBColor(1, 186, 239)      # #01BAEF
    charcoal = RGBColor(43, 45, 66)   # #2B2D42
    muted_gray = RGBColor(128, 128, 128)
    
    # Title Page/Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("TÀI LIỆU CẤU TRÚC MÃ NGUỒN VÀ BẢN ĐỒ HỆ THỐNG BACKEND\n")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(48)
    sub_run = subtitle_p.add_run("Dự án HảiSản.vn (shop_sea) - Môn học SWP391\nKiến Trúc Domain-Driven Design (DDD) & Hexagonal Architecture")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = muted_gray

    doc.add_page_break()

    # SECTION 1: INTRODUCTION
    add_heading_styled(doc, "1. Kiến Trúc Tổng Quan (DDD & Hexagonal Architecture)", level=1, color=navy)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống backend của ").font.name = 'Calibri'
    run_bold = p.add_run("HảiSản.vn")
    run_bold.bold = True
    run_bold.font.color.rgb = navy
    p.add_run(" đã hoàn tất việc chuyển dịch toàn bộ kiến trúc từ mô hình MVC truyền thống sang kiến trúc hướng miền chuyên nghiệp ").font.name = 'Calibri'
    p.add_run("Domain-Driven Design (DDD) và Kiến trúc Lục giác (Hexagonal Architecture / Ports & Adapters) ").font.name = 'Calibri'
    p.add_run("nhằm cô lập nghiệp vụ cốt lõi, nâng cao khả năng mở rộng hệ thống và dễ dàng tích hợp các dịch vụ bên ngoài.").font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.add_run("Mô hình hoạt động dựa trên các nguyên tắc phân tầng chặt chẽ:").font.name = 'Calibri'
    
    bullets = [
        ("Domain Layer (Lớp miền): ", "Chứa thực thể thuần khiết (Entities), các đối tượng giá trị (Value Objects), và định nghĩa các giao diện lưu trữ (Ports / Interfaces). Lớp này hoàn toàn cô lập, không phụ thuộc vào bất kỳ framework hay thư viện database nào."),
        ("Application Layer (Lớp ứng dụng): ", "Chứa các Use Cases thực thi quy trình nghiệp vụ đơn lẻ, phối hợp các thực thể miền và cổng lưu trữ để hoàn thành một nhiệm vụ cụ thể."),
        ("Infrastructure Layer (Lớp hạ tầng): ", "Triển khai thực tế các cổng lưu trữ (Adapters / Repositories) sử dụng Mongoose/MongoDB, Redis Cache, các dịch vụ bên ngoài như Cloudinary CDN, Gmail SMTP, Groq LLM API."),
        ("Presentation Layer (Lớp giao diện/API): ", "Đón tiếp các REST Request thông qua Express HTTP Controllers, thực thi validate dữ liệu đầu vào và trả về mã trạng thái JSON chuẩn hóa.")
    ]
    
    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_t = bp.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = navy
        bp.add_run(desc).font.name = 'Calibri'

    # SECTION 2: ROOT DIRECTORY STRUCTURE
    add_heading_styled(doc, "2. Bản Đồ Cấu Trúc Thư Mục Thực Tế", level=1, color=navy)
    
    table_p = doc.add_paragraph()
    table_p.add_run("Dưới đây là chi tiết các thành phần thư mục chính trong `backend/src` sau khi tái cấu trúc:").font.name = 'Calibri'
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Đường dẫn thư mục/tệp tin'
    hdr_cells[1].text = 'Phân lớp / Kiến trúc'
    hdr_cells[2].text = 'Chức năng chính trong hệ thống'
    
    # Format Header Row
    for cell in hdr_cells:
        set_cell_background(cell, "0B4F6C")
        set_cell_margins(cell, 120, 120, 150, 150)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Table rows
    rows_data = [
        ("src/shared/domain/", "Shared Kernel (DDD)", "Chứa các lớp trừu tượng dùng chung: Entity, ValueObject, AggregateRoot và hệ thống quản lý sự kiện miền DomainEvents."),
        ("src/modules/iam/", "Bounded Context (IAM)", "Quản lý tài khoản, phân quyền (User/Admin), xác thực JWT Stateless kết hợp Blacklist Redis, Google OAuth và OTP khôi phục mật khẩu."),
        ("src/modules/product/", "Bounded Context (Core)", "Quản lý mẻ hàng hải sản, phân loại Tươi/Khô, tích hợp định vị GeoJSON Point phục vụ truy vấn khoảng cách $near mặt cầu."),
        ("src/modules/post/", "Bounded Context (Forum)", "Điều phối diễn đàn cộng đồng, bình luận đa cấp, lượt thích (likes) realtime giữa các thành viên."),
        ("src/modules/recipe/", "Bounded Context (Recipes)", "Quản lý cẩm nang công thức chế biến ẩm thực hải sản, thời gian chuẩn bị và định lượng khẩu phần."),
        ("src/modules/boat-log/", "Bounded Context (Cabin)", "Nhật ký hành trình cabin đi biển thực tế dành riêng cho ngư dân uy tín chứng thực nguồn gốc hải sản."),
        ("src/repositories/", "Anti-Corruption Layer", "Lớp chống tham nhũng: Bọc các repo cũ (user.repository, product.repository, v.v.) thành adapter ánh xạ sang DDD Use Cases phục vụ tương thích ngược."),
        ("src/routes/ & src/controllers/", "Presentation / API Layer", "Chứa các endpoint Express, rate limiters, cookies parser và điều chuyển request sang lớp Presentation của DDD."),
        ("src/middlewares/", "Infrastructure / Core", "Middleware kiểm soát an toàn hệ thống: authenticate (JWT HttpOnly cookie), validate (Zod schema validation), upload (Multer memory buffering), csrf (Double Submit Cookie)."),
        ("src/validations/", "Infrastructure / Core", "Định nghĩa các Zod schemas kiểm định tính hợp lệ của dữ liệu đầu vào phía Client gửi lên trước khi xử lý."),
        ("src/utils/", "Helpers", "Công thức toán học Haversine tính khoảng cách GPS, logger Winston tự động xoay vòng hàng ngày, cấu hình kết nối Redis/Mongoose/Cloudinary.")
    ]
    
    for folder, layer, desc in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = folder
        row_cells[1].text = layer
        row_cells[2].text = desc
        
        # Apply borders and formatting
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, 100, 100, 150, 150)
            if i == 0:
                set_cell_background(cell, "F0F4F8")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = charcoal
                    if i == 0 or i == 1:
                        run.font.bold = True

    # Column widths
    widths = [Inches(1.8), Inches(1.8), Inches(3.4)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # SECTION 3: KEY ACHIEVEMENTS
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_heading_styled(doc, "3. Các Tính Năng Đột Phá Đã Đạt Được (Tuần 1 - 6)", level=1, color=navy)
    
    achievements = [
        ("A. Tái Cấu Trúc Toàn Diện Sang DDD & Hexagonal Architecture",
         "Nhóm đã dịch chuyển toàn bộ mã nguồn cũ sang mô hình DDD sạch sẽ, chia tách rõ ràng 5 Bounded Contexts độc lập trong `src/modules/`. Thiết lập cơ chế CQRS tách biệt luồng ghi nghiệp vụ (Write Use Cases) và luồng đọc nhanh (Fast Read Queries) trực tiếp từ Mongoose với độ phản hồi trung bình dưới 80ms. Đồng thời bọc lớp Anti-Corruption Layer (ACL) tương thích ngược hoàn hảo."),
        ("B. Tích Hợp Chat Realtime & Các Tính Năng Tương Tác Nâng Cao",
         "Hệ thống chat 1-1 Socket.io đã được bổ sung các tính năng nâng cao cấp doanh nghiệp bao gồm: Thu hồi tin nhắn (Recall) đồng bộ realtime, Chỉnh sửa nội dung tin nhắn đã gửi (Edit), và Thả cảm xúc Emoji (Reaction). Toàn bộ tin nhắn chat đều được lọc khử độc ký tự thẻ HTML (chống XSS) và giới hạn tần suất gửi tin nhắn chống spam qua Redis."),
        ("C. Định Vị Bản Đồ Leaflet & Công Thức Haversine Thực Tế",
         "Mẻ hải sản tươi sống (Fresh) bắt buộc có định vị. Sử dụng chỉ mục `2dsphere` và toán tử địa lý `$near` / `$geoWithin` của MongoDB kết hợp công thức toán học lượng giác Haversine tính toán khoảng cách đường chim bay dạng mặt cầu trên Trái Đất, lọc nhanh sản phẩm xung quanh người mua trong bán kính 20km."),
        ("D. Tự Động Hóa Thanh Toán Sepay Webhook & Khôi Phục OTP Email",
         "Tích hợp Sepay Webhook bất đồng bộ tự động nâng cấp tài khoản ngư dân lên Premium khi quét mã VietQR thành công. Áp dụng cơ chế an toàn timing-attack (safeCompare) kiểm tra chữ ký Webhook và tự động cascade xóa sạch Token phiên cũ trên Redis bắt buộc đăng nhập lại. Luồng OTP quên mật khẩu được mã hóa và lưu giữ tạm thời trên Redis (TTL 5 phút)."),
        ("E. Trợ Lý Chatbot AI Hải Sản Hỗ Trợ Người Dùng",
         "Tích hợp mô hình ngôn ngữ lớn Llama 3.1 qua Groq Cloud API làm Trợ lý AI tư vấn cách sơ chế, bảo quản hải sản, gợi ý thực đơn món ngon và hướng dẫn các chức năng trực quan ngay trên giao diện Web."),
        ("F. Kiểm Thử Tự Động Jest & Tài Liệu Swagger UI",
         "Backend được bao bọc bởi bộ kiểm thử tự động toàn diện Jest (hơn 16 test suites, 68 ca kiểm thử đơn vị/tích hợp) đạt độ phủ mã nguồn tốt, kết hợp tài liệu Swagger UI tương tác trực tiếp tại `/api-docs` và CI/CD GitHub Actions chạy tự động mỗi lần commit.")
    ]
    
    for title, body in achievements:
        add_heading_styled(doc, title, level=2, color=navy)
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(8)
        run_b = bp.add_run(body)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = charcoal

    # SECTION 4: MIDTERM REVIEW & DEVELOPMENTS
    add_heading_styled(doc, "4. Nhiệm Vụ Review & Phát Triển Trọng Tâm (Tuần 6 - 8)", level=1, color=navy)
    
    p3 = doc.add_paragraph()
    p3.add_run("Dựa trên kết quả báo cáo giữa kỳ, nhóm xác định các nhiệm vụ kỹ thuật nâng cao tiếp theo. Để đảm bảo tiến độ, công việc được phân chia rõ rệt: Hai thành viên nòng cốt kĩ thuật phụ trách phát triển chính gánh vác khối lượng công việc CỰC NHIỀU, hai thành viên còn lại phụ trách bổ trợ khối lượng công việc CỰC ÍT.").font.name = 'Calibri'
    
    add_heading_styled(doc, "4.1 Công việc Core phát triển lõi kỹ thuật (Thuận & Đức - CỰC NHIỀU)", level=2, color=navy)
    
    core_tasks = [
        ("Triển khai Signaling WebRTC cho Video Call: ", "Xây dựng Signaling Server trao đổi SDP Offer/Answer và ICE Candidates thông qua Socket.IO backend. Cấu hình hệ thống STUN/TURN servers của Google để kết nối P2P vượt qua NAT/Firewall."),
        ("Quản lý trạng thái Cuộc gọi: ", "Thiết lập hệ thống kiểm soát và truyền tín hiệu trạng thái cuộc gọi trực tiếp (Đang gọi, Rung chuông, Bận, Đã cúp máy, Không liên lạc được) đồng bộ thời gian thực."),
        ("Lập trình giao diện VideoCallOverlay: ", "Xây dựng overlay full-screen phía React client hỗ trợ truyền/nhận media stream từ camera/mic, xử lý bật/tắt thiết bị và kết nối kết thúc cuộc gọi."),
        ("Tối ưu hóa Database & Bảo mật: ", "Đánh chỉ mục Compound Indexes tối ưu hóa truy vấn tìm kiếm GPS cho hải sản tươi sống. Thực hiện rà soát lỗ hổng bảo mật chuyên sâu (NoSQL Injection, XSS sanitization) trên các route chat và sản phẩm."),
        ("Bộ kiểm thử tích hợp (Integration Tests): ", "Viết các ca kiểm thử tích hợp Jest cho toàn bộ luồng gọi điện WebRTC và cổng webhook Sepay để ngăn ngừa lỗi hồi quy.")
    ]
    
    for t_title, t_desc in core_tasks:
        ctp = doc.add_paragraph(style='List Bullet')
        ctp.paragraph_format.space_after = Pt(4)
        run_t = ctp.add_run(t_title)
        run_t.bold = True
        run_t.font.color.rgb = navy
        ctp.add_run(t_desc).font.name = 'Calibri'

    add_heading_styled(doc, "4.2 Công việc bổ trợ & kiểm thử (Bút & Cường - CỰC ÍT)", level=2, color=navy)
    
    support_tasks = [
        ("Cập nhật Tài liệu & Hướng dẫn: ", "Rà soát Swagger API Docs khớp với cấu trúc DDD mới, viết tài liệu README.md hướng dẫn cài đặt chạy thử local."),
        ("Seed Data & Slides: ", "Chuẩn bị nạp dữ liệu mẫu hải sản đa dạng vào MongoDB phục vụ demo, hỗ trợ làm slide thuyết trình và mô tả UI."),
        ("Micro-interactions & Responsive: ", "Tinh chỉnh CSS responsive nhỏ cho trang công thức trên màn hình di động, sửa đổi các hiệu ứng hover nút bấm giao diện."),
        ("Kiểm thử thủ công (Manual Testing): ", "Thực hiện test hộp đen (blackbox testing) gọi video WebRTC giữa hai máy khách khác nhau, phát hiện lỗi giật lag UI và thu thập phản hồi của người dùng.")
    ]
    
    for t_title, t_desc in support_tasks:
        stp = doc.add_paragraph(style='List Bullet')
        stp.paragraph_format.space_after = Pt(4)
        run_t = stp.add_run(t_title)
        run_t.bold = True
        run_t.font.color.rgb = navy
        stp.add_run(t_desc).font.name = 'Calibri'

    # Save document
    output_path = r"C:\Users\PC\OneDrive\Desktop\sea_shop\sea_shop\swp391-su26-ai-audit-project-swp391_se20a04_group-06\backend\system.docx"
    doc.save(output_path)
    print(f"Successfully generated Word document at {output_path}")

if __name__ == '__main__':
    main()
